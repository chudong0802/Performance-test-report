import matplotlib.pyplot as plt
import pandas as pd
import matplotlib.dates as mdate
from pylab import *
import numpy as np
import matplotlib.ticker as ticker
from docx import Document
from docx.shared import Inches
import os
#字体和日期格式的处理
plt.rcParams['font.sans-serif'] = ['SimHei']
dateparse = lambda dates: pd.datetime.strptime(dates,'%H:%M:%S')
name=pd.ExcelFile('analysis.xlsx')
sheetname=name.sheet_names
doc = Document()

def fun():
    for i in range(1,len(sheetname)):
        data = pd.read_excel('analysis.xlsx', sheet_name=sheetname[i], arse_dates=True, date_parser=dateparse)
        x=data['TIME']
        y1=data['DALVIK']
        y2=data['NATIVE']
        y3=data['CPU']
        fig = plt.figure(figsize=(15, 15))
        ax1=fig.add_subplot(311)
        ax1.plot(x,y1,label='DALVIK',linewidth=3.0,ms=10)
        ax1.set_ylabel('DALVIK',size=20,color='b')
        plt.setp(ax1.get_xticklabels(), fontsize=20,visible=False)
        plt.legend(loc=4,fontsize=20)

        ax2=fig.add_subplot(312,sharex=ax1)
        ax2.plot(x,y2,label='NATIVE',linewidth=3.0,ms=10,color='green')
        ax2.set_ylabel('NATIVE',size=20,color='b')
        plt.setp(ax2.get_xticklabels(), visible=False)
        plt.legend(loc=4,fontsize=20)


        ax3=fig.add_subplot(313,sharex=ax1)
        ax3.plot(x,y3,label='CPU',linewidth=3.0,ms=10,color='y')
        ax3.set_xlabel('TIME',size=20,color='b')
        ax3.set_ylabel('CPU',size=20,color='b')

        plt.legend(loc=0,fontsize=20,framealpha=0.2)
        plt.gca().xaxis.set_major_locator(ticker.MultipleLocator(60))
        plt.yticks(size=20)
        plt.xticks(size=20,rotation=60)
        plt.setp(ax3.get_xticklabels(), fontsize=20)
        plt.savefig('D:/untitled/HomeWork/Pic/TIME_DALVIK_NATIVE_CPU{}.png'.format(i))
        # dic=dict(zip(data['TIME'],data['DALVIK']))
        list=[]
        for j in range(len(data['TIME'])-1):
            if data['PID'][j]!=0:
                if data['PID'][j+1]!=data['PID'][j]:
                    list.append(data['TIME'][j + 1])
        # print(list)
        if list==[]:
            # print(sheetname[i])
            string = '{}没有出现crash'.format(sheetname[i])
            # images=open(imageFile)
            images = 'D:/untitled/HomeWork/Pic/TIME_DALVIK_NATIVE_CPU{}.png'.format(i)  # 保存在本地的图片
             # doc对象
            doc.add_paragraph(string)  # 添加文字
            doc.add_picture(images, height=Inches(5),width=Inches(5))# 添加图, 设置宽度
            doc.add_page_break()
            doc.save('word.docx')  # 保存路径
        else:
            # print(sheetname[i])
            string = '{}共crush{}次:分别在\n {}几个时间点'.format(sheetname[i], len(list), list)
            images = 'D:/untitled/HomeWork/Pic/TIME_DALVIK_NATIVE_CPU{}.png'.format(i)  # 保存在本地的图片
            doc.add_paragraph(string)  # 添加文字
            doc.add_picture(images,height=Inches(5),width=Inches(5)) # 添加图, 设置宽度
            doc.add_page_break()
            doc.save('word.docx')
if __name__ == '__main__':
    fun()



# #-*- coding:utf-8 -*-文字保留为图片
# from PIL import Image,ImageFont,ImageDraw
# dic=dict(zip(data['TIME'],data['DALVIK']))
# # print(dic.keys())
# # print(dic.values())
# list=[]
# for i in range(len(data['TIME'])-1):
#     if data['PID'][i+1]!=data['PID'][i]:
#         list.append(data['TIME'][i+1])
# print(list)
# text='共crush{}次:分别在\n {}几个时间点'.format(len(list),list)
# font = ImageFont.truetype("msyh.ttc",18)
# lines = []
# line =''
# for word in text.split():
#   print(word)
#   if font.getsize(line+word)[0] >= 300:
#     lines.append(line)
#     line = u''
#     line += word
#     # print('size=',font.getsize(line+word)[0])
#   else:
#     line = line + word
# line_height = font.getsize(text)[1]
# img_height = line_height*(len(lines)+1)
# print ('len=',len(lines))
# print ('lines=',lines)
# im = Image.new("RGB",(333,img_height),(255,255,255))
# dr = ImageDraw.Draw(im)
# x,y=5,5
# for line in lines:
#   dr.text((x,y),line,font=font,fill="#000000")
#   y += line_height
# im.save("1.1.jpg")



#将图片写入pdf文件
from reportlab.lib.pagesizes import portrait
from reportlab.pdfgen import canvas
from PIL import Image

# def jpg_to_pdf(jpg, pdf_path):
#     (w, h) = Image.open(jpg).size
#     user = canvas.Canvas(pdf_path, pagesize=portrait((w, h)))
#     user.drawImage(jpg, 0, 0, w, h)
#     user.showPage()
#     user.save()
# if __name__ == '__main__':
#     jpg_path = 'TIME_DALVIK_NATIVE_CPU.png'
#     pdf_path = 'code.pdf'
#     jpg_to_pdf(jpg_path, pdf_path)

# li=[]
# new_array=[]
# for label in range(len(list)):
#     x_label=list[label]
#     y_label=dic[x_label]
#
#     new_array.append([x_label,y_label])
# print(new_array)
# end_array=np.array(new_array)

# import  PyPDF2
# # 给指定的页面添加水印、公司标志或者时间戳。
# pdffile1 = open(r'code.pdf', 'rb')
# pdf_reader1 = PyPDF2.PdfFileReader(pdffile1)
# # 获取该文档的第一页
# first_page = pdf_reader1.getPage(0)
# # 打开印有水印的PDF文件
# water_pdf = PyPDF2.PdfFileReader(r'code1.pdf', 'rb')
# # 在上一个文档的第一页中加入这个有水印的文件
# first_page.mergePage(water_pdf.getPage(0))
# # 将读取的内容写入到对象中
# pdfwriter = PyPDF2.PdfFileWriter()
# pdfwriter.addPage(first_page)